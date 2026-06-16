# -*- coding: utf-8 -*-
"""VERTEXA — Şirkət & Maliyyə Planı sənədi (.docx). 2-fazalı strategiya."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import fabriq as F, model as FR

INK=RGBColor(0x0D,0x0E,0x10); AMBER=RGBColor(0xC9,0x8A,0x06); GREY=RGBColor(0x55,0x55,0x55)
GREEN=RGBColor(0x2E,0x7D,0x32); BLUE=RGBColor(0x1F,0x5F,0xA8)

rows,meta = F.compute(); r0,_ = F.compute(capital=0)
need = min(x['cum'] for x in r0); raise_amt = meta['capital']
be = next((x['m'] for x in rows if x['net']>0), None)
mincash = min(x['cum'] for x in rows)
def at(m): return rows[m-1]

doc = Document(); doc.styles['Normal'].font.name='Calibri'; doc.styles['Normal'].font.size=Pt(11)
def H(t,size=15,color=INK,after=6,before=12):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
def P(t,size=11,color=None,bold=False,italic=False,bullet=False,after=3):
    p=doc.add_paragraph(style='List Bullet' if bullet else None)
    r=p.add_run(t); r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color: r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(after)

# Title
t=doc.add_paragraph(); r=t.add_run("VERTEXA — Şirkət & Maliyyə Planı")
r.bold=True; r.font.size=Pt(22); r.font.color.rgb=INK
s=doc.add_paragraph(); rr=s.add_run("İstehsal-texnologiya şirkəti · 2-fazalı strategiya · ~$1M seed → öz gəliri ilə böyümə")
rr.italic=True; rr.font.size=Pt(12); rr.font.color.rgb=GREY

H("1. Vizyon")
P("VERTEXA — bir texnoloji nüvə (brauzer CAD + qiymət mühərriki + ödəniş + analitika) üzərində ardıcıl məhsullar çıxaran istehsal-texnologiya şirkətidir.")
for x in [
 "FEEDRATE — flaqman: anında qiymət → sifariş → supplier istehsalı → çatdırılma marketplace",
 "QUOTEFLOW — emalatxanalar üçün embed quote-widget SaaS",
 "FORMCHECK — DFM analiz SaaS (FAZA 2)",
 "CONFIGFLOW — e-commerce 3D konfiqurator (FAZA 2)",
]: P(x,bullet=True)
P("Nüvə ~70% hər məhsulda təkrar işlədilir → komanda ardıcıl axır, boş qalma yox.", italic=True, color=GREY)

H("2. 2-FAZALI STRATEGİYA (əsas)", color=BLUE)
P("FAZA 1 — Sübut (Ay 1–~20):", bold=True, color=AMBER)
for x in [
 "FEEDRATE işlək MVP Ay 1-də launch → gəlir 1-ci aydan (2 ay cilalanır)",
 "QUOTEFLOW Ay 5-də launch → 2 məhsul canlı",
 "12 nəfərlik komanda, ~$1M seed",
 f"Breakeven ~Ay {be} — 2 məhsul özünü sübut edir",
]: P(x,bullet=True)
P("FAZA 2 — Genişlənmə ÖZ GƏLİRİ ilə (Ay 22+):", bold=True, color=GREEN)
for x in [
 "FORMCHECK Ay 24-də, CONFIGFLOW Ay 30-da launch",
 "Yeni investisiya YOX — şirkət öz gəliri (breakeven sonrası kassa) ilə maliyyələşdirir",
 "Mövcud komanda məhsul 3-4-ü qurur (nüvə təkrar) — yalnız +2 dəstək (12 → 14 nəfər)",
]: P(x,bullet=True)
P(f"Kassa FAZA 2-də böyüyür: Ay 20 ~${at(20)['cum']:,.0f} → Ay 30 ~${at(30)['cum']:,.0f} → Ay 36 ~${at(36)['cum']:,.0f}. Yeni pul tələb olunmur.", bold=True, color=GREEN)

H("3. KOMANDA — FAZA 1 (12 nəfər)")
table=doc.add_table(rows=1,cols=4); table.style='Light Grid Accent 1'; table.alignment=WD_TABLE_ALIGNMENT.LEFT
for i,h in enumerate(["Rol","Başlama","Aylıq $","Qrup"]):
    c=table.rows[0].cells[i]; rp=c.paragraphs[0].add_run(h); rp.bold=True; rp.font.size=Pt(10)
phase1=[r for r in F.ROLES if r[1]<=2]
for role,sm,sal,heavy,grp in phase1:
    c=table.add_row().cells
    c[0].paragraphs[0].add_run(role).bold=True
    c[1].text=f"Ay {sm}"; c[2].text=f"{sal:,}"; c[3].text=grp
    for cell in c:
        for pp in cell.paragraphs:
            for rn in pp.runs: rn.font.size=Pt(9.5)
g1=sum(s for _,sm,s,*_ in phase1)
P(f"FAZA 1 komanda: {len(phase1)} nəfər · aylıq brüt ${g1:,} · +vergi(×1.25) ${g1*1.25:,.0f}/ay", bold=True, color=GREEN)
P("FAZA 2-də əlavə işçi: yalnız 2 müştəri dəstəyi. Məhsul 3-4-ü MÖVCUD developer komandası qurur (nüvə ~70% təkrar).", italic=True, color=GREY)

H("4. MALİYYƏ XÜLASƏSİ (USD)")
fin=[
 ("Tələb olunan seed (FAZA 1)", f"${raise_amt:,.0f}", "yalnız ilk 2 məhsul"),
 ("Real funding need (kapitalsız dib)", f"${-need:,.0f}", "ən dərin kassa"),
 ("Ən aşağı kassa (runway dibi)", f"${mincash:,.0f}", "müsbət qalır"),
 ("Breakeven ayı", f"Ay {be}", "ilk müsbət NET"),
 ("İlk 2 ay aylıq burn", f"${-(at(1)['net']+at(2)['net'])/2:,.0f}", "build dövrü"),
 ("24-cü ay aylıq gəlir", f"${at(24)['rev']:,.0f}", f"ARR ${at(24)['rev']*12:,.0f}"),
 ("48-ci ay aylıq gəlir", f"${at(48)['rev']:,.0f}", f"ARR ${at(48)['rev']*12:,.0f}"),
 ("48-ci ay aylıq NET", f"${at(48)['net']:,.0f}", "mənfəət"),
]
for name,val,note in fin:
    p=doc.add_paragraph()
    p.add_run(f"{name}: ").bold=True
    rr=p.add_run(val); rr.bold=True; rr.font.color.rgb=GREEN
    rn=p.add_run(f"  ({note})"); rn.italic=True; rn.font.size=Pt(9.5); rn.font.color.rgb=GREY
    p.paragraph_format.space_after=Pt(2)

H("5. NİYƏ BU GÜCLÜDÜR (investor üçün)")
for x in [
 "Az pul, az risk: ~$1M ilə 2 məhsul çıxır + breakeven — $2M+ tələb edən planlardan səmərəli",
 "Sübut-sonra-genişlən: investor 4 məhsula yox, sübut olunmuş 2 məhsula və davam planına investisiya edir",
 "Öz gəliri ilə böyümə: FAZA 2 yeni kapital tələb etmir → durulaşma az",
 "Paylaşılan nüvə: hər yeni məhsulun marjinal xərci aşağı (~70% kod təkrar)",
 "Erkən gəlir: FEEDRATE Ay 1-dən gəlir gətirir (işlək MVP)",
]: P(x,bullet=True,color=(GREEN if x.startswith("Öz") else None))

P("Valyuta USD (AZN üçün ×1.70). Bütün rəqəmlər maliyyə modelindən (VERTEXA-Maliyye-Model.xlsx) avtomatik gəlir.", italic=True, color=GREY, after=2)

doc.save("VERTEXA-Plan.docx")
print("OK docx saved")
