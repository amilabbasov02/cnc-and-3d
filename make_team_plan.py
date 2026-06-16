# -*- coding: utf-8 -*-
"""FEEDRATE — 2 aylıq intensiv start: komanda & rollar planı (.docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

INK = RGBColor(0x0D,0x0E,0x10)
AMBER = RGBColor(0xC9,0x8A,0x06)
GREY = RGBColor(0x55,0x55,0x55)
GREEN = RGBColor(0x2E,0x7D,0x32)
RED = RGBColor(0xB0,0x2A,0x2A)

doc = Document()
# base style
st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(11)

def H(text, size=16, color=INK, after=6, before=12):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    return p

def P(text, size=11, color=None, bold=False, italic=False, bullet=False, after=3):
    p = doc.add_paragraph(style='List Bullet' if bullet else None)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    return p

# ---------- TITLE ----------
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("FEEDRATE — 2 AYLIQ İNTENSİV START PLANI")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = INK
s = doc.add_paragraph(); rr = s.add_run("Komanda · Rollar · Paralel iş axınları · Kim nə edir")
rr.italic = True; rr.font.size = Pt(12); rr.font.color.rgb = GREY
P("Məqsəd: layihəni 5 ayda yox, 2 ayda cilalı və işlək launch etmək. Hər kritik rolda 2+ nəfər, paralel iş axınları, fasiləsiz test.", italic=True, color=AMBER, after=8)

# ---------- 1. PRINSIPLER ----------
H("1. Əsas prinsiplər")
for x in [
 "Paralel iş axınları — frontend, backend, CAD eyni anda işləyir, gözləmə yoxdur",
 "Hər kritik rolda 2+ nəfər — tək nəfər bottleneck/risk olmasın (biri xəstələnsə iş dayanmasın)",
 "Fasiləsiz test — QA paralel işləyir, hər gün test edir, bug dərhal düzəlir",
 "4 sprint × 2 həftə = 2 ay. Gündəlik standup (15 dəq), həftəlik demo",
 "Scope intizamı: «mükəmməl» = cilalı MVP deməkdir, hər funksiya yox. Nice-to-have-lar v2-yə",
 "Bir mənbə həqiqət: ortaq backlog/board (hər kəs nə üzərində işlədiyini görür)",
]:
    P(x, bullet=True)

# ---------- 2. KOMANDA CƏDVƏLİ ----------
H("2. Komanda — rol, sayı, səviyyə, nə edir")
# (rol, sayı, səviyyə, aylıq brüt $/nəfər, nə edir)
team = [
 ("CEO / Product lead", "1", "—", 2000, "Qərarlar, prioritet, koordinasiya, supplier & investor, gündəlik blokerləri açır"),
 ("CTO / Baş arxitekt", "1", "Senior", 4000, "Arxitektura, ən çətin nüvə, kod review, texniki qərarlar"),
 ("CAD developer (Three.js + CSG)", "2", "Senior", 3000, "Brauzer redaktor: primitivlər, boolean (deşik/cib), extrude, undo/redo, STL/OBJ export, anlıq qiymət geometriyası"),
 ("Frontend developer", "2", "Mid–Senior", 2200, "UI shell, model qaleriyası, redaktor interfeysi, sifariş/checkout axını, müştəri & supplier dashboard"),
 ("Backend developer", "2", "Mid–Senior", 2500, "Auth, DB (Postgres/Prisma), qiymət API, Stripe + escrow/payout, supplier marşrutlaşdırma, limit/anti-abuse"),
 ("UI/UX dizayner", "2", "Mid–Senior", 2000, "Dizayn sistemi, istifadəçi axınları, cilalama. (1 nəfər core məhsul, 1 nəfər marketplace/supplier tərəf)"),
 ("QA / Tester (ümumi app)", "2", "Mid", 1500, "Paralel funksional + regresiya test, cross-browser, mobil, bug triage, buraxılış nəzarəti"),
 ("CAD / Geometriya QA (Test)", "1", "Mid–Senior", 2000, "CAD redaktorun ÇIXIŞINI yoxlayır: boolean düzgünlüyü, mesh bütövlüyü (non-manifold/self-intersection), STL/OBJ/STEP export-un çap/emal üçün yararlılığı, ölçü dəqiqliyi, qiymət geometriyasının doğruluğu. (Mexaniki/CAD/3D arxa planı)"),
 ("Dizayn QA / UI-UX Reviewer", "1", "Mid–Senior", 1800, "DİZAYNI yoxlayır, təsdiqləyir: UI dizayn sisteminə uyğunluq (komponent/spacing/rəng/şrift), vizual ardıcıllıq, UX keyfiyyəti, responsiv & mobil görünüş, brend. Hər ekrana launch-dan əvvəl sign-off verir."),
 ("DevOps / İnfra", "1", "Senior", 2500, "CI/CD, staging+prod mühitlər, deploy, monitorinq, backup, təhlükəsizlik"),
 ("Operations / Supplier onboarding", "1", "—", 1500, "Supplier cəlb + yoxlama + capability, launch günü supply hazır olsun"),
 ("Marketinq / Launch", "1", "—", 2000, "Landing, brend, reklam hazır, launch kampaniyası 1-ci gün üçün hazır"),
 ("HR / İşə qəbul (Recruiter)", "1", "—", 1500, "18 nəfərlik komandanı TEZ işə götürmək: namizəd axtarışı, müsahibə, onboarding, müqavilələr, komanda məsələləri. HR olmadan bu tempdə işçi yığmaq mümkün deyil."),
]
table = doc.add_table(rows=1, cols=5); table.style = 'Light Grid Accent 1'; table.alignment = WD_TABLE_ALIGNMENT.LEFT
hdr = table.rows[0].cells
for i, h in enumerate(["Rol", "Sayı", "Səviyyə", "Aylıq $/nəfər", "Nə edir"]):
    hdr[i].text = ""; rp = hdr[i].paragraphs[0].add_run(h); rp.bold = True; rp.font.size = Pt(10)
for role, cnt, lvl, sal, what in team:
    c = table.add_row().cells
    c[0].paragraphs[0].add_run(role).bold = True
    c[1].text = cnt; c[2].text = lvl; c[3].text = f"{sal:,}"; c[4].text = what
    for cell in c:
        for pp in cell.paragraphs:
            for rn in pp.runs: rn.font.size = Pt(9.5)
# total
total = sum(int(x[1]) for x in team)
P(f"CƏMİ: {total} nəfər (2 aylıq build komandası). «Hər roldan 2+» + 3 nəzarət rolu (CAD QA, Dizayn QA, ümumi QA) + HR (komandanı yığmaq üçün).", bold=True, color=GREEN, after=8)

# ---------- BUDGET ----------
TAX = 0.25
monthly_gross = sum(int(c)*s for _,c,_,s,_ in team)
monthly_loaded = monthly_gross*(1+TAX)
build_payroll = monthly_loaded*2
dev_heavy = ["CTO","CAD developer","Frontend","Backend","DevOps","CAD / Geometriya"]
eq_heavy = sum(int(c) for r,c,_,_,_ in team if any(k in r for k in dev_heavy))
eq_office = total - eq_heavy
equipment = eq_heavy*2200 + eq_office*1300
infra_tools = 6000  # 2 ay hosting+alət+lisenziya
build_total = build_payroll + equipment + infra_tools
H("2.b. 2 AYLIQ BUILD BÜDCƏSİ (bu komanda üçün)", size=13, color=AMBER)
for label, val, note in [
 ("Aylıq brüt maaş (cəmi)", monthly_gross, f"{total} nəfər"),
 ("Aylıq maaş + vergi/SSF (×1.25)", monthly_loaded, "işəgötürən faktiki xərci"),
 ("2 aylıq maaş (build)", build_payroll, "build dövrü"),
 ("Avadanlıq (birdəfəlik)", equipment, f"{eq_heavy} dev iş stansiyası + {eq_office} ofis"),
 ("İnfra + alət (2 ay)", infra_tools, "hosting, lisenziya, alətlər"),
 ("CƏMİ — 2 AYLIQ BUILD", build_total, "launch-a qədər lazım olan"),
]:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: "); r1.bold = (label.startswith('CƏMİ'))
    r2 = p.add_run(f"${val:,.0f}"); r2.bold = True; r2.font.color.rgb = (GREEN if label.startswith('CƏMİ') else INK)
    r3 = p.add_run(f"  ({note})"); r3.italic = True; r3.font.size = Pt(9.5); r3.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(2)
P("Qeyd: bu, yalnız 2 aylıq BUILD xərcidir (gəlir-xərc proqnozu deyil). Launch-dan sonrakı aylar ayrıca hesablanır. Valyuta USD (AZN üçün ×1.70).", italic=True, color=GREY, after=6)

# ---------- 3. PARALEL İŞ AXINLARI ----------
H("3. Paralel iş axınları (kim eyni anda nə üzərində)")
streams = [
 ("A — CAD nüvə", "CTO + 2 CAD developer", "Brauzer redaktor, geometriya, qiymət mühərrikinin geometriya hissəsi"),
 ("B — Frontend / UI", "2 Frontend + 2 UX/UI + Dizayn QA", "Bütün ekranlar, axınlar, cilalama (core + supplier tərəf paralel) — Dizayn QA hər ekranı təsdiqləyir"),
 ("C — Backend / ödəniş", "2 Backend", "Auth, DB, qiymət API, Stripe/escrow, supplier routing"),
 ("D — Test", "2 ümumi QA + 1 CAD/Geometriya QA", "Hər sprint-də paralel test, regresiya, bug triage — CAD QA geometriya/export-u ayrıca yoxlayır"),
 ("E — Supply + Launch", "Ops + Marketinq", "Supplier onboarding + launch hazırlığı (məhsuldan asılı olmadan paralel)"),
 ("F — İnfra", "DevOps", "Mühitlər, deploy, monitorinq — komandanı sürətli saxlayır"),
 ("G — İşə qəbul / HR", "HR", "Komandanı yığmaq, müsahibə, onboarding, müqavilələr — 1-ci gündən paralel (komanda olmadan iş başlamaz)"),
]
for name, who, what in streams:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(f"{name}: "); r1.bold = True; r1.font.color.rgb = AMBER
    r2 = p.add_run(f"[{who}] — {what}"); r2.font.size = Pt(10.5)
P("Bütün axınlar EYNİ ANDA işləyir. CAD bitməsini gözləmədən frontend UI qurur, backend API hazırlayır, QA hazır hissələri test edir.", italic=True, color=GREY, after=6)

# ---------- 4. 2 AYLIQ TIMELINE ----------
H("4. 2 aylıq timeline (4 sprint × 2 həftə)")
sprints = [
 ("Sprint 1 — Həftə 1-2: Təməl & skelet", [
   "CTO/CAD: arxitektura + CAD redaktor skeleti (primitivlər, move/rotate/scale)",
   "Backend: Auth + DB sxem + qiymət API skeleti",
   "Frontend: UI shell + qaleriya + əsas naviqasiya",
   "UX: dizayn sistemi + əsas ekran maketləri",
   "DevOps: CI/CD + staging/prod mühitlər", "QA: test planı + mühit",
   "Ops/Marketinq: ilk supplier danışıqları + landing",
 ]),
 ("Sprint 2 — Həftə 3-4: Əsas funksiyalar", [
   "CAD: boolean (deşik/cib), extrude, undo/redo, STL/OBJ export",
   "Backend: Stripe inteqrasiya + sifariş axını + limit/anti-abuse",
   "Frontend: redaktor UI + anlıq qiymət ekranı + sifariş/checkout",
   "UX: cilalama + marketplace/supplier UX",
   "QA: funksional test başlayır (paralel)", "Ops: supplier onboarding + capability",
 ]),
 ("Sprint 3 — Həftə 5-6: Marketplace + supplier + escrow", [
   "Backend: supplier marşrutlaşdırma + escrow/payout + statement",
   "Frontend: supplier paneli + sifariş izləmə + müştəri dashboard",
   "CAD: performans + cilalama (STEP araşdırma — opsional, v2)",
   "UX: edge-case axınlar + son cilalama",
   "QA: inteqrasiya test + regresiya + cross-browser", "Ops: 20-30 supplier hazır",
 ]),
 ("Sprint 4 — Həftə 7-8: Cilalama + test + LAUNCH 🚀", [
   "Hamı: bug fix, performans, cilalama (polish)",
   "QA: tam regresiya + yük testi + təhlükəsizlik yoxlaması",
   "DevOps: prod hardening + monitorinq + backup",
   "Marketinq: launch kampaniyası canlı", "Ops: supplier-lər aktiv, ilk sifarişlər",
   "→ LAUNCH (2-ci ayın sonu)",
 ]),
]
for title, items in sprints:
    H(title, size=12, color=INK, after=2, before=8)
    for it in items:
        P(it, bullet=True, size=10.5)

# ---------- 5. RİSK / SCOPE QEYDİ ----------
H("5. Vacib qeyd — 2 ay aqressivdir, scope-u qoru")
P("2 ayda «mükəmməl» = CİLALI MVP deməkdir, hər şey yox. Bu tempi tutmaq üçün:", bold=True)
for x in [
 "Hər modulun yalnız ƏSAS dəyərini qur; «nice-to-have» (fillet, STEP export, mürəkkəb marketplace funksiyaları) v2-yə",
 "Scope dəyişikliyinə «yox» de — sprint ortasında yeni funksiya əlavə etmə",
 "Hər kritik rolda 2 nəfər olması məhz buna görədir: paralel + risksiz (biri düşsə iş dayanmır)",
 "Əgər 2 ay çox sıxışdırsa: launch-u «əsas axın mükəmməl» ilə et, ikinci dərəcəli modullar launch-dan sonra",
]:
    P(x, bullet=True, color=RED if x.startswith("Scope") else None)

# ---------- 6. QEYDLƏR ----------
H("6. Mənim qeydlərim (doldur)")
P("(Bu bölmə sənin üçün boşdur — hər rol/sprint üçün öz qeydlərini, adları, dəyişiklikləri bura yaz.)", italic=True, color=GREY)
for _ in range(8):
    doc.add_paragraph("…………………………………………………………………………………………………………")

doc.save("FEEDRATE-2Ay-Komanda-Plan-v2.docx")
print("OK docx saved")
