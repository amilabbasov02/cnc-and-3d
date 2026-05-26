# -*- coding: utf-8 -*-
"""FEEDRATE — yekun 2 sened: (1) Qlobal Muqayise, (2) Plan."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT=RGBColor(0xC9,0x8A,0x06); DARK=RGBColor(0x1A,0x1F,0x26)
GREEN=RGBColor(0x16,0xA3,0x4A); RED=RGBColor(0xC0,0x2A,0x2A)
GREY=RGBColor(0x5A,0x63,0x6E); BLUE=RGBColor(0x1D,0x4E,0xD8)

def setup(doc, landscape=False):
    st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(10.5)
    if landscape:
        s=doc.sections[0]; s.orientation=WD_ORIENT.LANDSCAPE
        s.page_width, s.page_height = Inches(11.69), Inches(8.27)
        s.left_margin=s.right_margin=Inches(0.6)

def shade(cell,c):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),c); tcPr.append(sh)
def H(doc,t,l=1,color=DARK):
    p=doc.add_heading(t,level=l)
    for r in p.runs: r.font.color.rgb=color
    return p
def P(doc,t,size=10.5,color=None,bold=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    r.bold=bold; return p
def BL(doc,items,mark="•",mc=DARK):
    for it in items:
        x=doc.add_paragraph(style="List Bullet")
        if mark!="•":
            r=x.add_run(mark+" "); r.bold=True; r.font.color.rgb=mc
        x.add_run(it)
def NUM(doc,items):
    for it in items: doc.add_paragraph(style="List Number").add_run(it)
def TBL(doc,headers,rows,fs=8.7,color_marks=False):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Light Grid Accent 1"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hc=t.rows[0].cells
    for i,ht in enumerate(headers):
        hc[i].text=""; run=hc[i].paragraphs[0].add_run(ht); run.bold=True; run.font.size=Pt(fs)
        run.font.color.rgb=RGBColor(255,255,255); shade(hc[i],"2A2E35")
    for row in rows:
        c=t.add_row().cells
        for i,v in enumerate(row):
            c[i].text=""; run=c[i].paragraphs[0].add_run(v); run.font.size=Pt(fs)
            if i==0: run.bold=True
            if color_marks:
                if v=="✓": run.font.color.rgb=GREEN
                elif v=="✗": run.font.color.rgb=RED
                elif v=="~": run.font.color.rgb=ACCENT
            if headers[i]=="FEEDRATE": run.bold=True
    return t
def title(doc,t,sub):
    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=tp.add_run(t); r.bold=True; r.font.size=Pt(19); r.font.color.rgb=DARK
    sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=sp.add_run(sub); r.font.size=Pt(11); r.font.color.rgb=ACCENT
    dp=doc.add_paragraph(); dp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=dp.add_run("Hazırlayan: Kayzen · Tarix: 24.05.2026"); r.font.size=Pt(9); r.font.color.rgb=GREY

# =====================================================================
# DOC 1 — QLOBAL MUQAYISE
# =====================================================================
d=Document(); setup(d, landscape=True)
title(d,"FEEDRATE — Qlobal Bazar və Rəqabət Müqayisəsi",
      "Kim necə edir · Supplier-li vs supplier-siz model · Sistem nə qədər qabağa gedə bilər")

H(d,"1. İcmal və əsas nəticə",1)
P(d,"Azərbaycan daxili bazarı bu sahə üçün kiçikdir → strategiya QLOBAL olmalıdır (İngilis dilli, Stripe, "
    "onlayn marketinq). Bazarda hər BİR funksiyanı (ani qiymət, brauzer CAD, model kitabxanası, dizayn "
    "satışı) edən var, amma HAMISINI bir self-serve əlçatan platformada birləşdirən YOXDUR — fərqlənmə budur.")
BL(d,[
 "Supplier-SİZ model (yalnız software) ilə başlamaq tövsiyə olunur — asset-light, tez gəlir, aşağı risk.",
 "Supplier-Lİ model (Xometry kimi) çox kapital/əməliyyat tələb edir; lider belə hələ zərərdədir — sonraya.",
 "Praktiki hədəf: əvvəl niş/regional-qlobal liderlik, sonra genişlənmə.",
])

H(d,"2. Bazar həcmi və trendlər",1)
P(d,"Rəqəmsal istehsal bazarı 2026-da ~$439 mlrd, illik ~16% (CAGR); on-demand seqment ~53%. Trendlər: "
    "AI/ML ani qiymət, anında DFM, CAD daxilində qiymət (add-in), Çin birbaşa-istehsal qiymət təzyiqi.")

H(d,"3. Bazar seqmentasiyası",1)
TBL(d,["Seqment","Kimə","Nümunələr"],[
 ["1. Marketplace / istehsal platforması","Alıcı","Xometry, Protolabs, Fictiv, RapidDirect, Haizol, JLCCNC"],
 ["2. Quoting SaaS (emalatxana/maker aləti)","İstifadəçi","Paperless Parts, DigiFabster, PartPilot, aShop"],
 ["3. ERP + quoting","Böyük emalatxana","E2, JobBOSS², ProShop, Global Shop"],
 ["4. Should-cost / maya","OEM/satınalma","aPriori, DFMA, Costimator"],
 ["5. Texnologiya təminatçısı","Qurucular","CAD Exchanger, HOOPS, OpenCASCADE, Datakit"],
])

H(d,"4. Əsas oyunçular — +/−",1)
players=[
 ("Xometry / Hubs","Ani qiymət marketplace (ML, DFM). +: güclü, geniş, qlobal. −: 20–40% marja, brend itkisi, asılılıq."),
 ("Protolabs","Öz fabriki, ən sürətli. +: sürət, etibar. −: bahalı, alət deyil (xidmət)."),
 ("Fictiv / RapidDirect","İdarə olunan/Çin marketplace + DFM. +: keyfiyyət/ucuz. −: marketplace, lojistika."),
 ("Çin: JLCCNC / PCBWay / Haizol","Birbaşa fabrik/aqreqator. +: 30–50% ucuz, ani (JLCCNC). −: keyfiyyət dəyişkən, dil/gömrük."),
 ("Paperless Parts","Emalatxana üçün dərin DFM+iş axını. +: güclü. −: çox bahalı, mürəkkəb, ABŞ."),
 ("DigiFabster","White-label quoting widget. +: öz brend, 24/7. −: $2k–$50k/il, 3D-print yönlü."),
 ("PartPilot / aShop","AI quoting (±5%). +: dəqiq, əlçatan. −: yeni, niş."),
 ("Onshape","Tam brauzer CAD (pro). +: güclü CAD. −: qiymət/marketplace yox, $1,500+/il."),
 ("eMachineShop","CAD + qiymət + istehsal (ən yaxın). +: birləşmə. −: köhnə/masaüstü CAD, dizayn satışı yox."),
 ("Cults3D / CGTrader / GrabCAD","Dizayn satış marketplace (20% kom.). +: satış. −: CNC qiymət/CAD yox, 3D-print/render."),
]
for n,desc in players:
    H(d,n,3,BLUE); P(d,desc)

H(d,"5. Bacarıq matrisi — kim hansını edir",1)
P(d,"İşarələr:  ✓ = var/güclü   ~ = qismən   ✗ = yox", size=9, color=GREY)
cols=["Bacarıq","Xometry","eMachineShop","Dassault Make","DigiFabster","Onshape","Brauzer CAD*","Model market.**","FEEDRATE"]
rows=[
 ["Ani CNC qiymət","✓","✓","✓","✓","✗","✗","✗","✓"],
 ["Hazır model kitabxanası","✗","~","~","✗","~","~","✓","✓"],
 ["Brauzerdə CAD (proqramsız)","✗","~","✓","✗","✓","✓","✗","✓"],
 ["Real DFM / feature","✓","~","✓","✓","~","✗","✗","✓"],
 ["Dizayn satış marketplace","✗","✗","✗","✗","✗","✗","✓","✓"],
 ["Self-serve freemium","~","~","✗","~","~","✓","✓","✓"],
 ["Əlçatan / öz platforma","✗","✗","✗","~","✗","✓","✓","✓"],
 ["HAMISI bir platformada","✗","~","~","✗","✗","✗","✗","✓"],
]
TBL(d,cols,rows,fs=8,color_marks=True)
P(d,"* Tinkercad/SelfCAD/Zoo.  ** Cults3D/CGTrader/GrabCAD.  Nəticə: yalnız FEEDRATE bütün sətirlərdə ✓.",
  size=8, color=ACCENT, bold=True)

H(d,"6. Ən yaxın rəqiblər və niyə yer boşdur",1)
BL(d,[
 "eMachineShop — CAD+qiymət+istehsal, amma CAD masaüstü/köhnə, dizayn satışı yox.",
 "Dassault 3DEXPERIENCE Make — CAD ekosistemi+qiymət, amma enterprise/bahalı/mürəkkəb.",
 "Onshape — güclü brauzer CAD, amma ani qiymət və dizayn satışı yox.",
])

H(d,"7. Supplier-Lİ vs Supplier-SİZ model müqayisəsi",1)
TBL(d,["Meyar","A — Supplier-siz","B — Supplier-li","C — Hibrid (A→B)"],[
 ["Gəlir","Abunə + dizayn komissiyası (5–10%)","Sifariş marjası 20–40%","Hər ikisi"],
 ["Kapital","Aşağı","Yüksək","Orta (mərhələli)"],
 ["Komanda","Kiçik (software)","Böyük (tədarük/QC/ops)","Mərhələli"],
 ["Gəlirə çatma","Tez (2–3 ay)","Yavaş","Tez başla, sonra genişlən"],
 ["Mənfəət","Real","Çox gec (lider zərərdə)","A-dan gəlir"],
 ["Risk","Konversiya","Toyuq-yumurta, QC, leakage","İcra ardıcıllığı"],
 ["Kiçik komanda üçün","Yüksək uyğun","Aşağı","Yüksək uyğun"],
])
P(d,"Reallıq: Xometry (lider) Q3 2025-də ~-$11.6M zərər; gəlirin ~89%-i təchizatçıya gedir. "
    "Two-sided marketplace toyuq-yumurta, keyfiyyət və leakage problemləri ilə doludur → B indi riskli.", color=RED)

H(d,"8. Nə qədər qabağa gedə bilər",1)
TBL(d,["Ssenari","Nəyə çatmaq","Realizm"],[
 ["Niş/regional-qlobal lider","Maker+SMB+dizayner+tələbə","Yüksək"],
 ["Qlobal əlçatan alternativ","Tanınmış brend","Orta"],
 ["Qlobal lider (Xometry səviyyə)","Bazar aparıcısı","Çətin (kapital/şəbəkə)"],
])

H(d,"9. Go/No-Go və verdikt",1,ACCENT)
TBL(d,["Meyar","Qiymət"],[
 ["Qlobal tələb","✓"],["Unikal birləşmə","✓"],["Texniki imkan (MVP var)","✓"],
 ["Vaxt (2–3 ay)","✓"],["Kapital — Variant A","✓"],["Kapital — Variant B","✗"],
 ["Qlobal marketinq","~"],["Rəqabət","~"],
], color_marks=True)
P(d,"Verdikt: BƏLİ — başla, amma Variant A (supplier-siz, qlobal, niş) ilə; B-ni traction+kapitaldan sonra. "
    "Tam komitmenti 90 günlük validasiyaya bağla (bax: Plan sənədi).", bold=True)

H(d,"10. Mənbələr",1)
for s in ["Xometry maliyyə — 3dprint.com, bowerycap.com","Marketplace çətinlikləri — sharetribe.com, stripe.com",
          "Bazar həcmi — businessresearchinsights.com, imarcgroup.com","Rəqiblər — rapiddirect.com, makerverse.com, selfcad.com, cults3d.com, emachineshop.com"]:
    x=d.add_paragraph(style="List Bullet"); r=x.add_run(s); r.font.size=Pt(9); r.font.color.rgb=GREY

d.save(r"c:\Users\SmartBee\Desktop\machine engineering\FEEDRATE-Qlobal-Muqayise.docx")
print("SAVED 1")

# =====================================================================
# DOC 2 — PLAN
# =====================================================================
p=Document(); setup(p, landscape=False)
title(p,"FEEDRATE — Məhsul, Texniki və Biznes Planı",
      "3D kitabxana · brauzer CAD (Three.js+CSG) · marketplace · freemium · yol xəritəsi")

H(p,"1. Vizyon",1)
P(p,"Qlobal, self-serve platforma: istifadəçi (1) hazır 3D model kitabxanasından seçib ölçü qoyub qiymət "
    "görür; (2) brauzerdə öz modelini qurur; (3) dizaynını satır. Başlanğıcda supplier YOXDUR — yalnız "
    "qiymət/dəyər + software gəliri (abunə + dizayn komissiyası).")

H(p,"2. İstifadəçi tipləri",1)
TBL(p,["Tip","Ehtiyac","Limit"],[
 ["Anonim","Tez qiymət","Cəmi 1 hesablama"],
 ["Pulsuz hesab","Müntəzəm istifadə","5 hesablama + 5 yükləmə / ay"],
 ["Ödənişli","Çox həcm/brend/API","Paketə görə"],
 ["Satıcı (designer)","Dizayn satmaq","Marketplace + payout"],
])

H(p,"3. İstifadəçi axınları",1)
P(p,"Anonim qiymət (1 dəfə) → qeydiyyat divarı → hesab (5/ay) → limit bitir → paket. "
    "CAD redaktorunda model qur → export (hesab+limit). Dizaynı listing kimi sat → ödəniş → payout.")

H(p,"4. Limit və anti-abuse",1)
TBL(p,["Səviyyə","Necə sayılır","Sıfırlanma","Qoruma"],[
 ["Anonim 1","Cookie + IP","Yox (qeydiyyata yönəldir)","Fingerprint"],
 ["Hesab 5/ay","Server sayğacı (user_id)","Aylıq","Email təsdiqi MƏCBURİ"],
 ["1 IP = 1 hesab","ip_registry yoxlaması","—","Email/SMS təsdiqi (VPN üçün)"],
])

H(p,"5. Paket / funksiya matrisi (təsdiqlənib)",1)
TBL(p,["Funksiya","Free $0","Starter $9","Pro $29","Business $99"],[
 ["Hesablama/ay","5","50","Limitsiz","Limitsiz"],
 ["Yükləmə/ay","5","25","100","Limitsiz"],
 ["Brauzer CAD","✓ məhdud","✓","✓","✓"],
 ["Brendli PDF","—","—","✓","✓"],
 ["Marketplace komissiyası (bizim xeyrimiz)","10%","8%","7%","5%"],
 ["API/komanda","—","—","—","✓"],
])

H(p,"6. Modullar",1)
H(p,"A) 3D model kitabxanası",2,BLUE)
P(p,"Qaleriya → seç → ölçü → qiymət. Mənbə: CC0/öz/generasiya/partnyor (yalnız icazəli — müəllif hüququ riski). "
    "Parametrik: bounding-box miqyas (MVP) və ya parametrik şablon. Həcm: Orta.")
H(p,"B) Brauzer CAD redaktoru — Three.js + CSG (təsdiqlənib)",2,BLUE)
P(p,"three + three-bvh-csg (boolean) + TransformControls. MVP (v1): primitivlər, move/rotate/scale, "
    "boolean (deşik/cib), extrude, undo/redo, export STL/OBJ, anlıq qiymət. v2: fillet/ölçü. v3: STEP export "
    "(B-Rep, OpenCascade.js). Data: əməliyyat tarixçəsi (feature tree, JSON). Həcm: Böyük (ən ağır hissə).")
P(p,"Qeyd: mesh əsaslıdır — peşəkar CAD-ı əvəz etmir, sadə parametrik detallar üçündür.", color=RED)
H(p,"C) Dizayn marketplace",2,BLUE)
P(p,"Listing (ad, qiymət, lisenziya, fayl) → alış (Stripe) → çatdırılma → komissiya + payout (Stripe Connect). "
    "Moderasiya + müəllif hüququ takedown prosesi. Həcm: Böyük.")
P(p,"Komissiya = hər satışdan 5–10% BİZİM xeyrimizdir (paketə görə dəyişir; abunə gəlirinin üstünə əlavə "
    "ikinci əsas gəlir mənbəyi). Aşağı komissiya çox satıcı cəlb edir — sonra tələbə görə qaldırıla bilər.",
    color=ACCENT, bold=True)

H(p,"7. Texniki arxitektura",1)
TBL(p,["Qat","Texnologiya"],[
 ["Frontend","Next.js + React + Three.js (mövcud)"],
 ["Auth","Auth.js (email/şifrə + təsdiq)"],
 ["Backend","Next.js route handlers"],
 ["Verilənlər bazası","PostgreSQL (Neon/Supabase) + Prisma"],
 ["Saxlama","S3 / Cloudflare R2"],
 ["Ödəniş","Stripe (Subscriptions + Connect)"],
 ["CAD nüvə","three + three-bvh-csg (client)"],
])
P(p,"DB cədvəllər: users, usage_monthly, ip_registry, models, designs, listings, orders, subscriptions.", size=9, color=GREY)

H(p,"8. Gəlir modeli və freemium iqtisadiyyatı",1)
P(p,"İki AYRI gəlir mənbəyi var — qarışdırmamaq vacibdir:")
TBL(p,["Gəlir növü","Kim kimə ödəyir","Bizə gələn"],[
 ["Abunə (paket)","İstifadəçi → biz","100% (yalnız Stripe ~3% bank haqqı çıxır)"],
 ["Marketplace satış","Alıcı → satıcı","Yalnız 5–10% komissiya bizim; qalan satıcıya"],
])
P(p,"Nümunə: $20-lıq dizayn satışında ~8% = $1.6 bizim, $18.4 satıcıya. Abunədə isə bütün məbləğ bizimdir.",
  color=GREY, size=9)
P(p,"Aşağıdakı cədvəl aylıq ümumi gəliri (abunə + komissiya) qiymətləndirir (fərziyyə):")
TBL(p,["Metrik","Konservativ","Orta","Optimist"],[
 ["Aylıq ziyarətçi","10,000","30,000","80,000"],
 ["Ödənişli istifadəçi","12","72","320"],
 ["Abunə gəliri/ay","~$288","~$1,944","~$9,600"],
 ["Marketplace GMV/ay (satış həcmi)","$1,000","$5,000","$20,000"],
 ["Komissiya gəliri (5–10%, ~8%)","~$80","~$400","~$1,600"],
 ["CƏMİ/ay","~$368","~$2,344","~$11,200"],
])

H(p,"9. KPI",1)
BL(p,["Aktivləşmə","Free→Paid konversiya (%)","MRR + churn","Aylıq hesablama/yükləmə","Marketplace GMV"])

H(p,"10. Yol xəritəsi + 90 günlük validasiya",1)
TBL(p,["Faza","Məzmun"],[
 ["Faza 1 — Təməl","Auth + DB + limit (anonim 1, hesab 5/ay, 1 IP=1) + mövcud quote/3D"],
 ["Faza 2 — Kitabxana","Lisenziyalı model qaleriyası + parametrik ölçü"],
 ["Faza 3 — Ödəniş","Stripe + paketlər"],
 ["Faza 4 — CAD MVP","three-bvh-csg redaktor + export limiti"],
 ["Faza 5 — Marketplace","Listing + Stripe Connect + payout"],
 ["Faza 6","CAD v2, STEP export, API"],
])
P(p,"90 günlük validasiya: 0–4 həftə MVP yayım; 4–8 həftə 500+ qeydiyyat; 8–12 həftə 10+ ödənişli VƏ/VƏ YA "
    "ilk dizayn satışı; Free→Paid ≥3–5% → GO (tam başla), əks halda pivot.", bold=True)

H(p,"11. Risklər",1)
TBL(p,["Risk","Yumşaltma"],[
 ["Modellərin müəllif hüququ","Yalnız icazəli/öz/generasiya + takedown"],
 ["IP-limit aldadılması","Email/SMS təsdiqi + rate-limit"],
 ["CAD redaktorunun həcmi","Mərhələli MVP (primitiv+boolean əvvəl)"],
 ["Ödəniş/payout hüquq","Stripe Connect (KYC/vergi)"],
 ["Aşağı konversiya","Limit balansı + dəyər + sürtünməsiz qeydiyyat"],
])

H(p,"12. Təsdiqlənmiş qərarlar",1)
BL(p,[
 "Paketlər: Free $0 / Starter $9 / Pro $29 / Business $99 — təsdiq.",
 "Marketplace komissiyası 5–10% (paketə görə) — abunədən sonra ikinci gəlir mənbəyimiz — təsdiq.",
 "CAD export MVP: yalnız STL/OBJ (STEP sonra) — təsdiq.",
 "Email təsdiqi məcburi — təsdiq.",
 "Başlanğıc model: supplier-SİZ (Variant A), qlobal — təsdiq.",
], "✓", GREEN)

p.save(r"c:\Users\SmartBee\Desktop\machine engineering\FEEDRATE-Plan.docx")
print("SAVED 2")
